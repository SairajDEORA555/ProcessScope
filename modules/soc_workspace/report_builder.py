"""Builds a Word (.docx) incident report from the TelemetryStore, in memory."""
import datetime
from io import BytesIO

from docx import Document

from core.telemetry import TelemetryStore

SEV_ORDER = {"Critical": 0, "High": 1, "Medium": 2}
PHASE_ORDER = {
    "Initial Access": 1, "Execution": 2, "Persistence": 3, "Privilege Escalation": 4,
    "Defense Evasion": 5, "Credential Access": 6, "Discovery": 7, "Lateral Movement": 8,
    "Command and Control": 9, "Exfiltration": 10, "Impact": 11,
}
IMPACT_BY_TACTIC = {
    "Impact": "Data destruction / ransomware - direct business disruption and potential data loss.",
    "Credential Access": "Credential compromise - risk of account takeover and lateral movement.",
    "Exfiltration": "Data exfiltration - possible breach with regulatory and reputational exposure.",
    "Command and Control": "Active attacker control of the host.",
    "Persistence": "Attacker maintains a foothold that survives reboots.",
    "Defense Evasion": "Evasion in use - detection gaps and possible tampering with tooling.",
    "Execution": "Arbitrary code execution achieved on the endpoint.",
}


def build_docx() -> bytes:
    alerts = sorted(TelemetryStore.alerts(),
                    key=lambda a: (SEV_ORDER.get(a["severity"], 9), a["time"]))
    doc = Document()
    doc.add_heading("ProcessScope Incident Report", 0)
    meta = doc.add_paragraph()
    meta.add_run("Generated: ").bold = True
    meta.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    if not alerts:
        doc.add_paragraph("No detections recorded. Run attacks in the Attack Simulator "
                          "to populate an incident.")
        return _save(doc)

    tactics = {a["tactic"] for a in alerts}
    severity = min((a["severity"] for a in alerts), key=lambda s: SEV_ORDER.get(s, 9))
    inc_id = f"INC-2026-{7000 + len(alerts)}"
    events = sorted(TelemetryStore.events(), key=lambda e: e["time"])
    window = f"{events[0]['time']} - {events[-1]['time']}" if events else "-"

    doc.add_heading("Incident Summary", level=1)
    s = doc.add_paragraph()
    for label, value in [("Incident ID", inc_id), ("Severity", severity),
                         ("Host", "DESKTOP-PS (user Manas)"), ("Time window", window),
                         ("Detections", str(len(alerts))), ("ATT&CK tactics", str(len(tactics)))]:
        s.add_run(f"{label}: ").bold = True
        s.add_run(f"{value}\n")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"A multi-stage intrusion was detected on host DESKTOP-PS involving {len(alerts)} "
        f"correlated detections across {len(tactics)} MITRE ATT&CK tactics. The highest observed "
        f"severity was {severity}. This report summarises the timeline, attack progression, "
        f"indicators of compromise, business impact, and recommended containment and recovery actions."
    )

    doc.add_heading("Detections", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Time", "Detection", "Severity", "Technique"
    for a in alerts:
        row = t.add_row().cells
        row[0].text, row[1].text = a["time"], a["name"]
        row[2].text, row[3].text = a["severity"], a["technique"]

    doc.add_heading("Timeline", level=1)
    tt = doc.add_table(rows=1, cols=3)
    tt.style = "Table Grid"
    h = tt.rows[0].cells
    h[0].text, h[1].text, h[2].text = "Time", "Source / Event", "Message"
    for e in events:
        row = tt.add_row().cells
        row[0].text = e["time"]
        row[1].text = f"{e['source']} {e['event_id']}"
        row[2].text = e["message"]

    doc.add_heading("Attack Story", level=1)
    ordered = sorted(alerts, key=lambda a: PHASE_ORDER.get(a["tactic"], 99))
    for i, a in enumerate(ordered, 1):
        doc.add_paragraph(f"{a['tactic']} - {a['name']} ({a['technique']}): {a['description']}",
                          style="List Number")

    doc.add_heading("MITRE ATT&CK Coverage", level=1)
    for tactic in sorted(tactics, key=lambda x: PHASE_ORDER.get(x, 99)):
        techs = ", ".join(sorted({a["technique"] for a in alerts if a["tactic"] == tactic}))
        doc.add_paragraph(f"{tactic}: {techs}", style="List Bullet")

    doc.add_heading("Indicators of Compromise", level=1)
    for ioc in sorted({i for a in alerts for i in a.get("iocs", [])}):
        doc.add_paragraph(ioc, style="List Bullet")

    doc.add_heading("Business Impact", level=1)
    rating = {"Critical": "Critical", "High": "High", "Medium": "Moderate"}.get(severity, "Moderate")
    doc.add_paragraph(f"Overall business impact: {rating}.")
    for tactic in sorted(tactics, key=lambda x: PHASE_ORDER.get(x, 99)):
        if tactic in IMPACT_BY_TACTIC:
            doc.add_paragraph(f"{tactic}: {IMPACT_BY_TACTIC[tactic]}", style="List Bullet")

    doc.add_heading("Containment & Recovery", level=1)
    for step in ["Isolate affected host(s) from the network",
                 "Terminate malicious processes and block C2 IPs / domains",
                 "Disable and rotate compromised credentials (and krbtgt if domain-wide)",
                 "Remove persistence mechanisms",
                 "Preserve forensic evidence (memory + disk image)",
                 "Validate offline backups and rebuild/restore affected systems",
                 "Apply hardening and monitor for recurrence"]:
        doc.add_paragraph(step, style="List Bullet")

    return _save(doc)


def _save(doc) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
